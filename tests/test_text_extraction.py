"""
tests/test_text_extraction.py
------------------------------
Unit tests for utils/text_extraction.py
No external dependencies required.
"""
from __future__ import annotations

import io
import sys
import os

# Make backend importable from the tests directory
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))

import pytest

from utils.text_extraction import extract_text, _extract_txt


class TestExtractTxt:
    def test_utf8_text(self):
        content = "Hello, world!\nThis is a test."
        result = extract_text(content.encode("utf-8"), "essay.txt")
        assert result == content

    def test_latin1_fallback(self):
        content = "Caf\xe9 au lait"
        result = extract_text(content.encode("latin-1"), "notes.txt")
        assert "Caf" in result

    def test_empty_file(self):
        result = extract_text(b"", "empty.txt")
        assert result == ""


class TestUnsupportedFormat:
    def test_raises_on_csv(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text(b"col1,col2", "data.csv")

    def test_raises_on_no_extension(self):
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract_text(b"some content", "README")


class TestDocxExtraction:
    def test_docx_extraction(self):
        """Build a real minimal .docx in memory and verify extraction."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("First paragraph of test document.")
        doc.add_paragraph("Second paragraph with more content.")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        result = extract_text(buf.read(), "test.docx")
        assert "First paragraph" in result
        assert "Second paragraph" in result
